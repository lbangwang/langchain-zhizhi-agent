"""职责：文档加载——纯文本 / Markdown / 简易 PDF / Word(.docx) 抽成纯文本。

技术点：pypdf 文本层（无 OCR）；python-docx；多编码尝试解码。
"""

from __future__ import annotations

import io
from pathlib import Path


SUPPORTED_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".csv",
    ".json",
    ".pdf",
    ".docx",
}


def _decode_bytes(raw: bytes) -> str:
    """功能：按 utf-8/gbk/latin-1 等编码尝试把字节解成字符串。

    技术点：多编码回退；最后 utf-8 errors=ignore。
    """
    for enc in ("utf-8", "utf-8-sig", "gbk", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def extract_pdf_text(raw: bytes) -> str:
    """功能：用 pypdf 抽取 PDF 可复制文本层。

    技术点：无 OCR；抽不出文本则抛 ValueError。
    """
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    parts: list[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:  # noqa: BLE001
            t = ""
        if t.strip():
            parts.append(t.strip())
    text = "\n\n".join(parts).strip()
    if not text:
        raise ValueError("PDF 未能抽出文本（可能是扫描件，请先 OCR 或粘贴文本）")
    return text


def extract_docx_text(raw: bytes) -> str:
    """功能：用 python-docx 抽取 Word 段落与表格文字。

    技术点：段落 + 表格单元格用 | 拼接。
    """
    from docx import Document

    doc = Document(io.BytesIO(raw))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n\n".join(parts).strip()
    if not text:
        raise ValueError("Word 文档无有效文本")
    return text


def extract_text_from_bytes(raw: bytes, filename: str) -> str:
    """功能：按扩展名解析上传字节为纯文本。

    技术点：pdf/docx 走专用抽取；.doc 拒绝；其余按文本解码。
    """
    if not raw:
        raise ValueError("空文件")
    suffix = Path(filename or "upload.txt").suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(raw)
    if suffix == ".docx":
        return extract_docx_text(raw)
    if suffix == ".doc":
        raise ValueError("暂不支持旧版 .doc，请另存为 .docx，或使用「粘贴文本」")
    if suffix and suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"不支持的格式 {suffix}，请使用 "
            f"{', '.join(sorted(SUPPORTED_EXTENSIONS))} 或粘贴文本"
        )
    return _decode_bytes(raw).strip()
