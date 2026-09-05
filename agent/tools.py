"""职责：Agent 核心工具——搜索 / 写 txt、PDF、Word(.docx)。

技术点：LangChain @tool；写盘前 HITL；产物登记 artifact；ContextVar 带 user/chat。
"""

from __future__ import annotations

import contextvars
import re
from io import BytesIO
from pathlib import Path

import httpx
from langchain_core.tools import tool

from app.config import get_settings
from app.db import SessionLocal
from app.models import Artifact, ToolAudit
from app.stop_signal import is_stopped
from app.utils import new_id, utcnow
from agent.hitl import require_hitl_or_skip

# 运行上下文：由 runner 在调用前设置
current_user_id: contextvars.ContextVar[str] = contextvars.ContextVar(
    "agent_user_id", default=""
)
current_chat_id: contextvars.ContextVar[str] = contextvars.ContextVar(
    "agent_chat_id", default=""
)
current_config_version: contextvars.ContextVar[str] = contextvars.ContextVar(
    "agent_config_version", default=""
)

# 优先 TTF（fpdf2 对 TTC 支持不稳定）；Arial Unicode 含中日韩字形
_CJK_FONT_CANDIDATES = [
    Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
    Path("/Library/Fonts/Arial Unicode.ttf"),
    Path("/System/Library/Fonts/Supplemental/Songti.ttc"),
    Path("/System/Library/Fonts/STHeiti Light.ttc"),
    Path("/System/Library/Fonts/Hiragino Sans GB.ttc"),
    Path("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc"),
    Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
]

DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _audit(tool_name: str, inp: str, out: str, status: str = "ok") -> None:
    """功能：把一次工具调用写入 tool_audit 表。

    技术点：ContextVar 取 user_id/chat_id/config_version；无库则跳过。
    """
    if SessionLocal is None:
        return
    uid = current_user_id.get() or "unknown"
    with SessionLocal() as db:
        db.add(
            ToolAudit(
                id=new_id(),
                user_id=uid,
                chat_id=current_chat_id.get() or None,
                tool_name=tool_name,
                input_preview=(inp or "")[:500],
                output_preview=(out or "")[:1000],
                status=status,
                config_version=current_config_version.get() or None,
                create_date=utcnow(),
            )
        )
        db.commit()


def _save_artifact(filename: str, data: bytes, content_type: str) -> str:
    """功能：保存产物文件并登记 artifact 表，返回 artifact_id。

    技术点：落盘 artifacts_dir；软删字段 is_del=0。
    """
    settings = get_settings()
    root = Path(settings.artifacts_dir)
    root.mkdir(parents=True, exist_ok=True)
    artifact_id = new_id()
    safe_name = filename.replace("/", "_")
    path = root / f"{artifact_id}_{safe_name}"
    path.write_bytes(data)
    if SessionLocal is None:
        return artifact_id
    uid = current_user_id.get() or "unknown"
    with SessionLocal() as db:
        db.add(
            Artifact(
                id=artifact_id,
                user_id=uid,
                chat_id=current_chat_id.get() or None,
                filename=safe_name,
                content_type=content_type,
                storage_path=str(path),
                byte_size=len(data),
                create_date=utcnow(),
                is_del=0,
            )
        )
        db.commit()
    return artifact_id


def _safe_filename(title: str, ext: str) -> str:
    """功能：从标题生成短文件名，去掉【多Agent】等前缀避免叠两层。

    技术点：正则清洗非法路径字符；无有效字时回退 report。
    """
    q = title or ""
    for _ in range(4):
        nxt = re.sub(r"【[^】]*】", "", q)
        if nxt == q:
            break
        q = nxt
    q = re.split(r"[，,。；;]|采用|导出", q, maxsplit=1)[0]
    cleaned = re.sub(r"[\\/:*?\"<>|\s【】\[\]]+", "_", q.strip())
    cleaned = re.sub(r"_+", "_", cleaned).strip("._")[:24]
    if not cleaned or not re.search(r"[A-Za-z0-9\u4e00-\u9fff]", cleaned):
        cleaned = "report"
    ext = (ext or "txt").lstrip(".").lower()
    return f"{cleaned}.{ext}"


def infer_export_format(task: str) -> str:
    """功能：从用户任务判断交付格式：txt / pdf / docx。

    技术点：关键词优先 pdf；doc/文档映射 docx，不生成旧版 .doc。
    """
    t = (task or "").lower()
    if "pdf" in t:
        return "pdf"
    if any(k in t for k in ("docx", "word", "doc文档", "word文档")):
        return "docx"
    if re.search(r"(^|[^a-z])doc([^a-z]|$)", t) or "文档" in (task or ""):
        return "docx"
    if any(k in t for k in ("txt", "文本", ".md", "markdown")):
        return "txt"
    return "pdf"


def _resolve_cjk_font() -> Path | None:
    """功能：返回本机可用的中文字体文件；没有则 None（PDF 将降级 ASCII）。

    技术点：按候选路径探测 TTF/TTC；fpdf2 对 TTC 支持不稳定故优先 TTF。
    """
    for p in _CJK_FONT_CANDIDATES:
        if p.is_file():
            return p
    return None


def _build_txt_bytes(title: str, body: str) -> bytes:
    """功能：把标题和正文编码成 UTF-8 文本字节。

    技术点：纯文本 fallback，供 PDF/Word 失败时改落 txt。
    """
    return f"{title or '报告'}\n\n{(body or '').strip()}\n".encode("utf-8")


def _build_pdf_bytes(title: str, body: str) -> bytes:
    """功能：生成 PDF 字节；中文依赖本机 CJK 字体。

    技术点：fpdf2；无 CJK 字体时 ASCII replace 降级。
    """
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    font_path = _resolve_cjk_font()
    use_cjk = False
    if font_path is not None:
        try:
            pdf.add_font("CJK", fname=str(font_path))
            pdf.set_font("CJK", size=14)
            use_cjk = True
        except Exception:  # noqa: BLE001
            use_cjk = False
    if not use_cjk:
        pdf.set_font("Helvetica", size=14)
        title = (title or "").encode("ascii", "replace").decode("ascii")
        body = (body or "").encode("ascii", "replace").decode("ascii")

    pdf.multi_cell(0, 8, title or "Report")
    pdf.ln(3)
    pdf.set_font("CJK" if use_cjk else "Helvetica", size=11)
    text = (body or "").strip()[:12000] or "(empty)"
    pdf.multi_cell(0, 6, text)
    raw = pdf.output()
    return bytes(raw) if isinstance(raw, (bytes, bytearray)) else str(raw).encode()


def _build_docx_bytes(title: str, body: str) -> bytes:
    """功能：生成 Word .docx 字节（兼容用户口头说的 doc）。

    技术点：python-docx；eastAsia 宋体，避免中文乱码。
    """
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "宋体")

    doc.add_heading(title or "报告", level=1)
    for line in (body or "").splitlines() or [""]:
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _save_export(filename: str, data: bytes, content_type: str, kind: str = "") -> str:
    """功能：落盘产物并返回给用户看的短文案（不含内部路径）。

    技术点：委托 _save_artifact；文案指向前端「产物」下载。
    """
    _save_artifact(filename, data, content_type)
    labels = {"txt": "文本", "pdf": "PDF", "docx": "Word 文档", "doc": "Word 文档"}
    label = labels.get(kind or Path(filename).suffix.lstrip("."), "文件")
    return f"已生成{label}「{filename}」，可在右侧「产物」下载。"


def invoke_export(fmt: str, title: str, body: str) -> str:
    """功能：按格式调用对应写文件工具（含 HITL），供多 Agent / 超级智能体兜底。

    技术点：LangChain tool.invoke；txt/docx/pdf 三路分发。
    """
    kind = (fmt or "pdf").lower()
    if kind == "txt":
        return write_text_file.invoke(
            {
                "filename": _safe_filename(title, "txt"),
                "content": f"{title}\n\n{body}",
            }
        )
    if kind in {"doc", "docx", "word"}:
        return create_doc_report.invoke({"title": title, "body": body})
    return create_pdf_report.invoke({"title": title, "body": body})


@tool
def search_web(query: str) -> str:
    """功能：搜索公开网页信息（有 TAVILY_API_KEY 时用 Tavily，否则返回占位结果）。

    技术点：LangChain @tool；Tavily HTTP；is_stopped 则跳过。
    """
    chat_id = current_chat_id.get()
    if chat_id and is_stopped(chat_id):
        return "任务已停止，跳过搜索"
    settings = get_settings()
    try:
        if settings.tavily_api_key:
            resp = httpx.post(
                "https://api.tavily.com/search",
                json={
                    "api_key": settings.tavily_api_key,
                    "query": query,
                    "max_results": 3,
                },
                timeout=20,
            )
            resp.raise_for_status()
            data = resp.json()
            lines = []
            for item in data.get("results", [])[:3]:
                lines.append(
                    f"- {item.get('title')}: {item.get('content', '')[:180]} ({item.get('url')})"
                )
            out = "\n".join(lines) or "无搜索结果"
        else:
            out = (
                f"（未配置 TAVILY_API_KEY，占位搜索）关于「{query}」："
                "建议配置 Tavily 以获得真实网页结果。"
            )
        _audit("search_web", query, out)
        return out
    except Exception as exc:  # noqa: BLE001
        msg = f"搜索失败: {exc}"
        _audit("search_web", query, msg, status="error")
        return msg


@tool
def write_text_file(filename: str, content: str) -> str:
    """功能：将文本写入产物文件（可下载）。filename 仅文件名，不要包含路径。

    技术点：HITL require_hitl_or_skip；artifact 落盘；停止信号跳过。
    """
    chat_id = current_chat_id.get()
    if chat_id and is_stopped(chat_id):
        return "任务已停止，跳过写文件"
    denied = require_hitl_or_skip(
        "write_text_file", f"{filename}:{(content or '')[:80]}"
    )
    if denied:
        _audit("write_text_file", filename, denied, status="rejected")
        return denied
    name = Path(filename).name or "output.txt"
    if not name.endswith((".txt", ".md", ".json", ".csv", ".log")):
        name = f"{name}.txt"
    try:
        data = content.encode("utf-8")
        _save_artifact(name, data, "text/plain; charset=utf-8")
        out = f"已写入文本「{name}」，可在右侧「产物」下载。"
        _audit("write_text_file", f"{name}:{content[:80]}", out)
        return out
    except Exception as exc:  # noqa: BLE001
        msg = f"写文件失败: {exc}"
        _audit("write_text_file", filename, msg, status="error")
        return msg


@tool
def create_pdf_report(title: str, body: str) -> str:
    """功能：生成 PDF 报告并登记为可下载产物。用户要求「生成 PDF / 导出报告」时必须调用本工具。

    技术点：HITL；fpdf2；失败依次回退 docx、txt。
    """
    chat_id = current_chat_id.get()
    if chat_id and is_stopped(chat_id):
        return "任务已停止，跳过 PDF"
    denied = require_hitl_or_skip("create_pdf_report", f"{title}:{(body or '')[:80]}")
    if denied:
        _audit("create_pdf_report", title, denied, status="rejected")
        return denied
    try:
        data = _build_pdf_bytes(title, body)
        fname = _safe_filename(title, "pdf")
        out = _save_export(fname, data, "application/pdf", kind="pdf")
        _audit("create_pdf_report", title, out)
        return out
    except Exception as exc:  # noqa: BLE001
        # PDF 失败时改落 Word，避免再变成「用户要 doc 却只给 txt」
        try:
            data = _build_docx_bytes(title, body)
            fname = _safe_filename(title, "docx")
            saved = _save_export(fname, data, DOCX_MIME, kind="docx")
            msg = f"PDF 未生成成功，已改为 Word 文档。{saved}"
            _audit("create_pdf_report", title, msg, status="fallback")
            return msg
        except Exception as exc2:  # noqa: BLE001
            fname = _safe_filename(title, "txt")
            _save_artifact(
                fname,
                _build_txt_bytes(title, body),
                "text/plain; charset=utf-8",
            )
            msg = f"PDF/Word 未生成成功，已改为文本「{fname}」，可在右侧「产物」下载。"
            _audit("create_pdf_report", title, msg, status="fallback")
            return msg


@tool
def create_doc_report(title: str, body: str) -> str:
    """功能：生成 Word 文档（.docx）。用户要求 doc / docx / Word / 文档 时调用本工具。

    技术点：HITL；python-docx；失败回退 txt。旧版二进制 .doc 不生成。
    """
    chat_id = current_chat_id.get()
    if chat_id and is_stopped(chat_id):
        return "任务已停止，跳过 Word 文档"
    denied = require_hitl_or_skip("create_doc_report", f"{title}:{(body or '')[:80]}")
    if denied:
        _audit("create_doc_report", title, denied, status="rejected")
        return denied
    try:
        data = _build_docx_bytes(title, body)
        fname = _safe_filename(title, "docx")
        out = _save_export(fname, data, DOCX_MIME, kind="docx")
        _audit("create_doc_report", title, out)
        return out
    except Exception as exc:  # noqa: BLE001
        fname = _safe_filename(title, "txt")
        _save_artifact(
            fname,
            _build_txt_bytes(title, body),
            "text/plain; charset=utf-8",
        )
        msg = f"Word 未生成成功，已改为文本「{fname}」，可在右侧「产物」下载。"
        _audit("create_doc_report", title, msg, status="fallback")
        return msg


@tool
def search_images(query: str) -> str:
    """功能：通过 MCP 图片搜索服务检索相关图片链接（需 MCP_ENABLED）。

    技术点：LangChain @tool；检索类不走 HITL；委托 mcp_client。
    """
    chat_id = current_chat_id.get()
    if chat_id and is_stopped(chat_id):
        return "任务已停止，跳过图片搜索"
    # 检索类工具不走 HITL，避免多步任务被反复打断
    settings = get_settings()
    if not settings.mcp_enabled:
        out = f"（MCP 未启用）占位结果：关于「{query}」可搜索公开图库。"
        _audit("search_images", query, out)
        return out
    try:
        from agent.mcp_client import search_images_via_mcp

        out = search_images_via_mcp(query)
        _audit("search_images", query, out)
        return out
    except Exception as exc:  # noqa: BLE001
        msg = f"图片搜索失败: {exc}"
        _audit("search_images", query, msg, status="error")
        return msg


@tool
def load_skill(skill_id: str) -> str:
    """功能：按目录索引中的 skill_id 加载完整 Skill 流程说明（按需，省 system token）。

    技术点：委托 skills_loader.load_skill_body；检索类不走 HITL；未知 id 提示可选列表。
    """
    chat_id = current_chat_id.get()
    if chat_id and is_stopped(chat_id):
        return "任务已停止，跳过加载 Skill"
    from agent.skills_loader import list_skill_ids, load_skill_body

    sid = (skill_id or "").strip()
    body = load_skill_body(sid)
    if not body:
        known = ", ".join(list_skill_ids()) or "（当前无已安装 Skill）"
        out = f"未知 skill_id={sid!r}。可选：{known}"
        _audit("load_skill", sid, out, status="error")
        return out
    out = f"# Skill `{sid}`\n\n{body}"
    _audit("load_skill", sid, out[:500])
    return out


CORE_TOOLS = [
    load_skill,
    search_web,
    write_text_file,
    create_pdf_report,
    create_doc_report,
    search_images,
]
